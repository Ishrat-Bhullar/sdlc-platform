"""
main.py — FastAPI app: CORS for the React frontend, cookie-based JWT auth,
and every endpoint needed for the Friday demo.

Run:
    uvicorn main:app --reload --port 8000

Endpoints
---------
Auth:
    POST /auth/register
    POST /auth/login        (sets access_token + refresh_token HttpOnly cookies)
    GET  /auth/me
    POST /auth/refresh       (bonus — mints a new access token from the refresh cookie)
    POST /auth/logout        (bonus — clears both cookies)

Project wizard & ingestion:
    POST /projects
    POST /ingestion/upload
    DELETE /projects/{project_id}

Live dashboard:
    GET  /dashboard/timeline

Workspace delivery:
    GET  /projects/{project_id}/deliverables
    GET  /generated_artifacts
"""
import os
import uuid
import logging
from datetime import datetime, timedelta, timezone
from pathlib import Path

import jwt
from cryptography.fernet import Fernet
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, Response, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from passlib.context import CryptContext
from pydantic import BaseModel
from sqlalchemy.orm import Session

from .models import (
    AgentRun,
    Base,
    DashboardTimelineResponse,
    Document,
    DocumentUploadResponse,
    GeneratedArtifact,
    GeneratedArtifactOut,
    LoginRequest,
    Project,
    ProjectCreate,
    ProjectDeliverable,
    ProjectDeliverablesResponse,
    ProjectOut,
    ProjectStatus,
    ProviderConfiguration,
    RunStatus,
    TimelineEvent,
    User,
    UserCreate,
    UserOut,
    engine,
    get_db,
)

# ===========================================================================
# Configuration
# ===========================================================================
JWT_SECRET_KEY = os.getenv("JWT_SECRET_KEY", "VfiW3bUtPut1O7yDWkPxujl6ua3FpjRx7P_QjQmmjtRel8OqGbR7IXZkyCjzmx7L")
JWT_REFRESH_SECRET_KEY = os.getenv("JWT_REFRESH_SECRET_KEY", "Kna7ilHIQH5cRA7CmzfcAxyEq6hQVzl_J2KRx9LJCVAmmKxJlunMMe-K-ld5oDeC")
JWT_ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "30"))
REFRESH_TOKEN_EXPIRE_DAYS = int(os.getenv("REFRESH_TOKEN_EXPIRE_DAYS", "7"))

PROVIDER_KEY_ENCRYPTION_KEY = os.getenv("PROVIDER_KEY_ENCRYPTION_KEY", "wSqu6lOQVJ2WhQddQB-TNdPSBQVmeLVC7AQ-9hszUDY=")
_fernet = Fernet(PROVIDER_KEY_ENCRYPTION_KEY.encode())

COOKIE_SECURE = os.getenv("COOKIE_SECURE", "false").lower() == "true"
STORAGE_BASE_PATH = Path(os.getenv("STORAGE_BASE_PATH", "./storage"))
DEMO_MODE = os.getenv("DEMO_MODE", "false").lower() == "true"
DEMO_EMAIL = "ishratbhullar@gmail.com"
logger = logging.getLogger("sdlc.demo")

CORS_ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.getenv(
        "FRONTEND_URLS",
        "http://localhost:5173,http://127.0.0.1:5173,http://localhost:3000,http://127.0.0.1:3000",
    ).split(",")
    if origin.strip()
]

ACCESS_COOKIE_NAME = "access_token"
REFRESH_COOKIE_NAME = "refresh_token"

def _extract_document_text(path) -> str:
    """Extract plain text from an uploaded BRD/RFP document (PDF or DOCX) so
    ingestion actually grounds the pipeline in the real document content.
    Falls back gracefully (empty string) for unsupported/unreadable files —
    upload still succeeds, it just won't enrich the project description."""
    from pathlib import Path as _Path
    p = _Path(path)
    suffix = p.suffix.lower()
    try:
        if suffix == ".pdf":
            import pdfplumber
            with pdfplumber.open(str(p)) as pdf:
                return "\n\n".join((page.extract_text() or "") for page in pdf.pages)
        if suffix in (".docx", ".doc"):
            import docx  # python-docx
            d = docx.Document(str(p))
            return "\n".join(para.text for para in d.paragraphs)
        if suffix in (".txt", ".md"):
            return p.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        logger.warning("[Ingestion] _extract_document_text failed for %s: %s", p.name, exc)
    return ""


pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")
# In DEMO_MODE, all paths bypass auth — the sentinel value "*" is checked below.
DEMO_AUTH_BYPASS_ALL = True  # set to False to require login even in demo

# ===========================================================================
# App + CORS
# ===========================================================================
app = FastAPI(title="EY Autonomous SDLC Studio")

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)


@app.on_event("startup")
def on_startup() -> None:
    Base.metadata.create_all(bind=engine)
    STORAGE_BASE_PATH.mkdir(parents=True, exist_ok=True)


@app.get("/health", tags=["meta"])
def health() -> dict:
    return {"status": "ok"}


# ===========================================================================
# Auth helpers
# ===========================================================================
def hash_password(plain_password: str) -> str:
    return pwd_context.hash(plain_password)


def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


def _create_token(user_id: int, token_type: str, secret: str, expires_delta: timedelta) -> str:
    now = datetime.now(timezone.utc)
    payload = {"sub": str(user_id), "type": token_type, "iat": now, "exp": now + expires_delta}
    return jwt.encode(payload, secret, algorithm=JWT_ALGORITHM)


def create_access_token(user_id: int) -> str:
    return _create_token(user_id, "access", JWT_SECRET_KEY, timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))


def create_refresh_token(user_id: int) -> str:
    return _create_token(user_id, "refresh", JWT_REFRESH_SECRET_KEY, timedelta(days=REFRESH_TOKEN_EXPIRE_DAYS))


def _decode_token(token: str, secret: str, expected_type: str) -> dict:
    try:
        payload = jwt.decode(token, secret, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Token has expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Invalid token")
    if payload.get("type") != expected_type:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, f"Expected a {expected_type} token")
    return payload


def _set_auth_cookies(response: Response, user_id: int) -> None:
    response.set_cookie(
        key=ACCESS_COOKIE_NAME,
        value=create_access_token(user_id),
        httponly=True,
        secure=COOKIE_SECURE,
        samesite="lax",
        max_age=ACCESS_TOKEN_EXPIRE_MINUTES * 60,
        path="/",
    )
    response.set_cookie(
        key=REFRESH_COOKIE_NAME,
        value=create_refresh_token(user_id),
        httponly=True,
        secure=COOKIE_SECURE,
        samesite="lax",
        max_age=REFRESH_TOKEN_EXPIRE_DAYS * 24 * 60 * 60,
        path="/",
    )


def _get_or_create_demo_user(db: Session) -> User:
    user = db.query(User).filter(User.email == DEMO_EMAIL).first()
    if user is None:
        user = User(
            email=DEMO_EMAIL,
            full_name="Demo User",
            role="developer",
            hashed_password=hash_password("DemoPassword123!"),
        )
        db.add(user)
        db.commit()
        db.refresh(user)
    return user


def get_current_user(request: Request, db: Session = Depends(get_db)) -> User:
    # In DEMO_MODE, bypass auth entirely for all paths
    if DEMO_MODE and DEMO_AUTH_BYPASS_ALL:
        token = request.cookies.get(ACCESS_COOKIE_NAME)
        if token:
            try:
                payload = _decode_token(token, JWT_SECRET_KEY, expected_type="access")
                user = db.get(User, int(payload["sub"]))
                if user:
                    return user
            except Exception:
                pass
        logger.info("DEMO_AUTH_BYPASS path=%s reason=demo_mode email=%s", request.url.path, DEMO_EMAIL)
        return _get_or_create_demo_user(db)

    token = request.cookies.get(ACCESS_COOKIE_NAME)
    if not token:
        logger.info("AUTH_REQUIRED path=%s demo_mode=%s reason=missing_cookie", request.url.path, DEMO_MODE)
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Not authenticated")
    try:
        payload = _decode_token(token, JWT_SECRET_KEY, expected_type="access")
    except HTTPException:
        raise
    user = db.get(User, int(payload["sub"]))
    if user is None:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "User no longer exists")
    return user


def encrypt_secret(plaintext: str) -> str:
    return _fernet.encrypt(plaintext.encode()).decode()


def _get_project_or_404(db: Session, project_id: int) -> Project:
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, f"Project {project_id} not found")
    return project


# ===========================================================================
# Auth endpoints
# ===========================================================================
@app.post("/auth/register", response_model=UserOut, status_code=status.HTTP_201_CREATED, tags=["auth"])
def register(payload: UserCreate, db: Session = Depends(get_db)) -> User:
    if db.query(User).filter(User.email == payload.email).first():
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "A user with that email already exists")
    user = User(
        email=payload.email,
        full_name=payload.full_name,
        role=payload.role.value,
        hashed_password=hash_password(payload.password),
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@app.post("/auth/login", response_model=UserOut, tags=["auth"])
def login(payload: LoginRequest, response: Response, db: Session = Depends(get_db)) -> User:
    if DEMO_MODE and payload.email.lower() == DEMO_EMAIL:
        user = _get_or_create_demo_user(db)
        _set_auth_cookies(response, user.id)
        return user

    user = db.query(User).filter(User.email == payload.email).first()
    if user is None or not verify_password(payload.password, user.hashed_password):
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Incorrect email or password")
    _set_auth_cookies(response, user.id)
    return user


@app.get("/auth/me", response_model=UserOut, tags=["auth"])
def me(current_user: User = Depends(get_current_user)) -> User:
    return current_user


@app.post("/auth/refresh", tags=["auth"])
def refresh(request: Request, response: Response, db: Session = Depends(get_db)) -> dict:
    token = request.cookies.get(REFRESH_COOKIE_NAME)
    if not token:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "No refresh token provided")
    payload = _decode_token(token, JWT_REFRESH_SECRET_KEY, expected_type="refresh")
    user_id = int(payload["sub"])
    if db.get(User, user_id) is None:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "User no longer exists")
    response.set_cookie(
        key=ACCESS_COOKIE_NAME,
        value=create_access_token(user_id),
        httponly=True,
        secure=COOKIE_SECURE,
        samesite="lax",
        max_age=ACCESS_TOKEN_EXPIRE_MINUTES * 60,
        path="/",
    )
    return {"detail": "Access token refreshed"}


@app.post("/auth/logout", tags=["auth"])
def logout(response: Response) -> dict:
    response.delete_cookie(ACCESS_COOKIE_NAME, path="/")
    response.delete_cookie(REFRESH_COOKIE_NAME, path="/")
    return {"detail": "Logged out"}


# ===========================================================================
# Project wizard & document ingestion
# ===========================================================================
@app.post("/projects", response_model=ProjectOut, status_code=status.HTTP_201_CREATED, tags=["projects"])
def create_project(
    payload: ProjectCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> Project:
    logger.info(
        "PROJECT_CREATE_REQUEST user=%s demo_mode=%s name=%s execution_mode=%s deliverables=%s providers=%s",
        current_user.email,
        DEMO_MODE,
        payload.project_name,
        payload.execution_mode.value,
        [deliverable.value for deliverable in payload.deliverables],
        [provider.value for provider in payload.providers.keys()],
    )
    project = Project(
        name=payload.project_name,
        description=payload.description,
        project_type=payload.project_type.value,
        execution_mode=payload.execution_mode.value,
        build_type=payload.build_type.value,
        status="in_progress",
    )
    db.add(project)
    db.flush()

    for deliverable_type in payload.deliverables:
        db.add(ProjectDeliverable(
            project_id=project.id,
            deliverable_type=deliverable_type.value,
            selected=True,
            status=RunStatus.PENDING.value,
        ))

    for provider_name, raw_key in payload.providers.items():
        db.add(ProviderConfiguration(
            project_id=project.id,
            provider_name=provider_name.value,
            enabled=True,
            encrypted_key=encrypt_secret(raw_key),
        ))

    db.add(TimelineEvent(project_id=project.id, stage="Project Created", status=RunStatus.COMPLETED.value))
    db.add(TimelineEvent(
        project_id=project.id,
        stage=f"Launch Mode: {payload.launch_mode or payload.execution_mode.value}",
        status=RunStatus.COMPLETED.value,
    ))
    if payload.manual_stages:
        db.add(TimelineEvent(
            project_id=project.id,
            stage=f"Manual Stages: {', '.join(payload.manual_stages)}",
            status=RunStatus.COMPLETED.value,
        ))
    if payload.build_profile:
        db.add(TimelineEvent(
            project_id=project.id,
            stage=f"Build Profile: {payload.build_profile}",
            status=RunStatus.COMPLETED.value,
        ))
    if payload.providers:
        db.add(TimelineEvent(
            project_id=project.id,
            stage=f"BYOK Providers: {', '.join(provider.value for provider in payload.providers.keys())}",
            status=RunStatus.COMPLETED.value,
        ))

    db.commit()
    db.refresh(project)
    logger.info("PROJECT_CREATE_SUCCESS project_id=%s user=%s status=201", project.id, current_user.email)
    return project


@app.post("/ingestion/upload", response_model=DocumentUploadResponse, tags=["ingestion"])
def upload_document(
    project_id: int = Form(..., description="Required so the document can be attached to a project"),
    document_type: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Project not found")

    project_dir = STORAGE_BASE_PATH / f"project_{project_id}"
    project_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    destination = project_dir / safe_name
    with destination.open("wb") as out_file:
        out_file.write(file.file.read())

    document = Document(
        project_id=project_id,
        file_name=file.filename,
        document_type=document_type,
        storage_path=str(destination),
    )
    db.add(document)
    db.add(TimelineEvent(
        project_id=project_id,
        stage=f"Document Uploaded: {file.filename}",
        status=RunStatus.COMPLETED.value,
    ))

    # Extract real document text (PDF/DOCX) and fold it into the project
    # description so the agent pipeline is actually grounded in the uploaded
    # BRD/RFP content, not just whatever free-text the user typed at creation.
    extracted_chars = 0
    try:
        extracted = _extract_document_text(destination)
        if extracted.strip():
            extracted_chars = len(extracted)
            existing = (project.description or "").strip()
            marker = f"\n\n--- Uploaded document: {file.filename} ---\n"
            project.description = (existing + marker + extracted[:12000]).strip()
            db.add(TimelineEvent(
                project_id=project_id,
                stage=f"Document Ingested: {file.filename} ({extracted_chars} chars extracted)",
                status=RunStatus.COMPLETED.value,
            ))
    except Exception as exc:
        logger.warning("[Ingestion] Text extraction failed for %s: %s", file.filename, exc)

    db.commit()
    db.refresh(document)

    return {"document_id": document.id, "upload_status": "success", "project_reference": project_id,
            "extracted_chars": extracted_chars}


# ===========================================================================
# Live dashboard timeline
# ===========================================================================
@app.get("/dashboard/timeline", response_model=DashboardTimelineResponse, tags=["dashboard"])
def get_dashboard_timeline(
    project_id: int | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = (
        db.get(Project, project_id)
        if project_id is not None
        else db.query(Project).order_by(Project.created_at.desc()).first()
    )
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "No project found")

    timeline_events = (
        db.query(TimelineEvent)
        .filter(TimelineEvent.project_id == project.id)
        .order_by(TimelineEvent.timestamp.asc())
        .all()
    )
    agent_runs = (
        db.query(AgentRun)
        .filter(AgentRun.project_id == project.id)
        .order_by(AgentRun.id.asc())
        .all()
    )

    return {
        "project_id": project.id,
        "project_status": project.status,
        "timeline_events": timeline_events,
        "agent_runs": agent_runs,
    }


# ===========================================================================
# Workspace delivery
# ===========================================================================
@app.get("/projects/{project_id}/deliverables", response_model=ProjectDeliverablesResponse, tags=["projects"])
def get_project_deliverables(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Project not found")

    deliverables = db.query(ProjectDeliverable).filter(ProjectDeliverable.project_id == project_id).all()
    artifacts = (
        db.query(GeneratedArtifact)
        .filter(GeneratedArtifact.project_id == project_id)
        .order_by(GeneratedArtifact.created_at.asc())
        .all()
    )
    return {"project_id": project_id, "deliverables": deliverables, "artifacts": artifacts}


@app.get("/generated_artifacts", response_model=list[GeneratedArtifactOut], tags=["projects"])
def list_generated_artifacts(
    project_id: int | None = None,
    artifact_type: str | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[GeneratedArtifact]:
    query = db.query(GeneratedArtifact)
    if project_id is not None:
        query = query.filter(GeneratedArtifact.project_id == project_id)
    if artifact_type is not None:
        query = query.filter(GeneratedArtifact.artifact_type == artifact_type)
    return query.order_by(GeneratedArtifact.created_at.asc()).all()


class BuildStartRequest(BaseModel):
    project_id: int


@app.post("/build/start", tags=["agents"])
async def start_build(
    payload: BuildStartRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = db.get(Project, payload.project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Project not found")

    logger.info(
        "BUILD_START_REQUEST project_id=%s user=%s demo_mode=%s",
        payload.project_id,
        current_user.email,
        DEMO_MODE,
    )

    from . import agent_runner

    runs = agent_runner.ensure_agent_runs_exist(db, payload.project_id)

    project.status = ProjectStatus.IN_PROGRESS.value
    db.add(TimelineEvent(project_id=payload.project_id, stage="Autonomous Build Started", status=RunStatus.RUNNING.value))
    if DEMO_MODE:
        db.add(TimelineEvent(
            project_id=payload.project_id,
            stage="Demo Build Queued",
            status=RunStatus.COMPLETED.value,
        ))
        db.commit()
        import asyncio
        asyncio.create_task(agent_runner.run_pipeline(payload.project_id))
        logger.info(
            "BUILD_START_SUCCESS project_id=%s agents_queued=%s demo_mode=%s status=200",
            payload.project_id,
            len(runs),
            DEMO_MODE,
        )
        return {
            "project_id": payload.project_id,
            "agents_queued": len(runs),
            "status": "started",
            "message": f"Demo autonomous build started for project {payload.project_id}",
        }

    db.commit()

    import asyncio
    asyncio.create_task(agent_runner.run_pipeline(payload.project_id))
    logger.info(
        "BUILD_START_SUCCESS project_id=%s agents_queued=%s demo_mode=%s status=200",
        payload.project_id,
        len(runs),
        DEMO_MODE,
    )

    return {
        "project_id": payload.project_id,
        "agents_queued": len(runs),
        "status": "started",
        "message": f"Autonomous build started for project {payload.project_id}",
    }


@app.delete("/projects/{project_id}", status_code=status.HTTP_200_OK, tags=["projects"])
def delete_project(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    """
    Delete a project and all related records.
    Cascading deletes handle agent_runs, approvals, artifacts, timeline_events,
    documents, deliverables, provider_configurations, and review_results.
    """
    project = _get_project_or_404(db, project_id)
    
    logger.info(
        "PROJECT_DELETE_REQUEST project_id=%s name=%s user=%s",
        project_id, project.name, current_user.email,
    )
    
    db.delete(project)
    db.commit()
    
    logger.info("PROJECT_DELETE_SUCCESS project_id=%s", project_id)
    
    return {
        "project_id": project_id,
        "status": "deleted",
        "message": f"Project '{project.name}' and all related records deleted successfully",
    }


from .main_extension import router as ext_router
app.include_router(ext_router)
from . import main_extension



if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)